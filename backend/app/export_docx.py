"""
Generate a matching Word (.docx) resume from structured data using
python-docx -- fully local, no external service.
"""
import io

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x7C, 0x1F, 0xD1)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)


def _contact_line(contact) -> str:
    parts = [p for p in [contact.email, contact.phone, contact.linkedin, contact.github, contact.location] if p]
    return "  |  ".join(parts)


def _section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT
    # thin rule under the heading
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "7C1FD1")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(2)


def build_resume_docx(resume) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(resume.contact.name or "Your Name")
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = DARK
    name_p.paragraph_format.space_after = Pt(2)

    contact_line = _contact_line(resume.contact)
    if contact_line:
        c_p = doc.add_paragraph()
        c_run = c_p.add_run(contact_line)
        c_run.font.size = Pt(9.5)
        c_run.font.color.rgb = GRAY
        c_p.paragraph_format.space_after = Pt(8)

    if resume.summary:
        _section_heading(doc, "Summary")
        p = doc.add_paragraph()
        run = p.add_run(resume.summary)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK

    if resume.skill_categories:
        _section_heading(doc, "Skills")
        for cat in resume.skill_categories:
            if not cat.skills:
                continue
            p = doc.add_paragraph()
            r1 = p.add_run(f"{cat.category}: ")
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = DARK
            r2 = p.add_run(", ".join(cat.skills))
            r2.font.size = Pt(10)
            r2.font.color.rgb = DARK
            p.paragraph_format.space_after = Pt(2)

    if resume.experience:
        _section_heading(doc, "Experience")
        for entry in resume.experience:
            header = entry.title or "Role"
            if entry.company:
                header += f" — {entry.company}"
            p = doc.add_paragraph()
            r = p.add_run(header)
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = DARK
            p.paragraph_format.space_after = Pt(1)
            if entry.dates:
                dp = doc.add_paragraph()
                dr = dp.add_run(entry.dates)
                dr.italic = True
                dr.font.size = Pt(9.5)
                dr.font.color.rgb = GRAY
                dp.paragraph_format.space_after = Pt(2)
            for b in entry.bullets:
                _bullet(doc, b)

    if resume.projects:
        _section_heading(doc, "Projects")
        for entry in resume.projects:
            p = doc.add_paragraph()
            r = p.add_run(entry.title)
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = DARK
            p.paragraph_format.space_after = Pt(1)
            for b in entry.bullets:
                _bullet(doc, b)

    if resume.education:
        _section_heading(doc, "Education")
        for entry in resume.education:
            header = entry.degree or "Degree"
            if entry.institution:
                header += f" — {entry.institution}"
            p = doc.add_paragraph()
            r = p.add_run(header)
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = DARK
            p.paragraph_format.space_after = Pt(1)
            if entry.dates:
                dp = doc.add_paragraph()
                dr = dp.add_run(entry.dates)
                dr.italic = True
                dr.font.size = Pt(9.5)
                dr.font.color.rgb = GRAY
                dp.paragraph_format.space_after = Pt(4)

    if resume.certifications:
        _section_heading(doc, "Certifications")
        for c in resume.certifications:
            _bullet(doc, c)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
